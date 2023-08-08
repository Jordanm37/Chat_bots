# from .openai_config import *
from dotenv import load_dotenv
import openai
import os 
load_dotenv(".env")
openai.api_key = os.environ.get("OPENAI_API_KEY")
openai.api_key = os.getenv('OPENAI_API_KEY')


import logging
logging.basicConfig(filename='example.log',level=logging.INFO)
logger = logging.getLogger(__name__)

logging.info('Started')
logging.info(f"OpenAI API key: {openai.api_key}")

from fastapi import FastAPI, File, UploadFile
from fastapi import Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware

import shutil
from docx import Document
from pydantic import BaseModel


# from .helpers import process_resume_file, process_job_link
# from .generator import generate_cover_letter


# Helpers 

from langchain.document_loaders import Docx2txtLoader
from docx import Document
from langchain.document_loaders import WebBaseLoader

def process_job_link(job_url):
    # Logging the URL to check if it's correct
    logger.info(f"Processing job URL: {job_url}")

    # Basic validation to check if it's a string and starts with "http"
    # if not isinstance(job_url, str) or not job_url.startswith("http"):
    if not isinstance(job_url, str):
        logging.error(f"Invalid job URL: {job_url}")
        return None

    url_loader = WebBaseLoader(job_url)
    postings = url_loader.load()
    # job_descr = [i.page_content for i in postings]
    job_description = postings[0].page_content

    return job_description


def process_resume_file(resume_file):
    
    resume_loader = Docx2txtLoader(resume_file)
    resume = resume_loader.load()[0].page_content
    
    return resume

def create_docx(text):
    doc = Document()

    file_path = "output.docx"
    # Split the cover letter text into paragraphs by new lines
    paragraphs = text.split('\n')

    # Iterate through the paragraphs and add them to the document
    for paragraph in paragraphs:
        if paragraph:  # Ignore empty lines
            doc.add_paragraph(paragraph)

    # Save the document
    doc.save(file_path)
    
    return file_path
    
def chatbot_completition(system_message, user_message, temp=None):
    if temp is not None:
        response = openai.ChatCompletion.create(
            model = 'gpt-4',
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
            temperature = temp,
        )
    else:
        response = openai.ChatCompletion.create(
            model = 'gpt-4',
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
            ],
        )    
    
    prompt, reply = response['usage']['prompt_tokens'], response['usage']['completion_tokens']
    cost = (0.03/1000)*prompt + (0.06/1000)*reply
    ## calculate api costs and log it
    logger.info(f"Prompt tokens: {prompt}, Completion tokens: {reply}, Cost: {cost}")

    return response["choices"][0]["message"]["content"]
def extract_key_words_and_resp(text):
    """
    Extracts key words from a text
    """
    SYS = """You are an expert career coach and hiring manager, responsible for extracting key aspects from a job posting. Your task includes identifying essential responsibilities, skills, and keywords. Do not introduce any additional information.

    The following examples should be excluded from your extraction as they do not represent intangible or technical skills, competencies, or job responsibilities:

    Location specific information (e.g., "Melbourne")
    Citizenship requirements (e.g., "Australian Citizen/Permanent Resident")
    Job perks and incentives (e.g., "Uncapped commission", "Weekly bonuses and commissions")
    Specific company details or affiliations (e.g., "Potential for a part of the global company", "Global blue-chip client", "American Express")
    Physical mobility and travel needs (e.g., "Having a current valid driver's license and a reliable vehicle", "Regional/interstate travel")
    Provided equipment (e.g., "Tools of the trade (tablet and phone) provided")
    Specific experience in certain industry other than sales or account management  (e.g., "3+ years sales experience within the international freight industry")
    Employment type (e.g., "Full time maximum term contract")
    Personal and career development opportunities (e.g., "Career change", "Entry level sales", "Development program", "Job training", "Mentorship", "Personalised coaching", "High achievers", "Future leaders")
    Superannuation details (e.g., "Super")
    Specific initiatives or programs (e.g., "Microcredentials initiatives", "Learning and Leadership Enterprise (LLE)")
    Instructions to review documentation (e.g., "Review the Position Description")
    Application process details (e.g., "Responses to the Selection Criteria and Core Competencies", "Online responses to pre-screening employment questions", "Application process understanding")
    Job application preparation (e.g., "Resume creation", "Cover letter writing")
    Support offered to candidates (e.g., "Support for candidates requiring adjustments", "Business School support")
    Privacy concerns (e.g., "Privacy of personal information")
    
    ie DONT NOT EXTRACT things that are not intagible skills/competencies or techincal skills or competencies.Your task is to focus on the technical and intangible skills, competencies, and job responsibilities that are described in the job posting. 
            
    Respond using markdown"""
    
    USER = """Given the following job posting, extract keywords that would be relevant to include in the experiences of a resume. Think this through step by step. Look for the relevant detail and keywords from the job posting:
                \n\n"""
    words = chatbot_completition(SYS, USER + text)
    
    
    return words


def filter_key_words(key_words, resume):
    """
    Extracts key words from a text
    """
    SYS = """You are an expert career coach and hiring manager, responsible for tailoring a resume perfectly for a job posting. Your task is to filter the given key words from a job posting  according to the resume you are provided with. Only remove key words that cannot be tailored to the resume or that the resume experience could not be tailored ot those key words.  Do not introduce any additional information.
            
    Format as a single list of key words. Respond using markdown"""
    
    USER = f"""Given the following keywords, extract those that would be relevant to include with the experiences of the provided resume. Think this through step by step.
    
    Key words to filter from job posting: {key_words}
    Resume: {resume}
                \n\n"""
    output = chatbot_completition(SYS, USER)
    
    
    return output
def extract_resume_exp(key_words, resume):
    """
    Extracts key words from a text
    """
    SYS = """You are an expert career coach and hiring manager, responsible for tailoring a resume perfectly for a job posting. Your task is to extract the experience from the given resume you are provided with that is related to the key words and job position provided. Only extract experience that can be tailored to the key words.  Do not introduce any additional information.
            
    Do not reduce the detail of the extracted experience. Respond using markdown"""
    
    USER = f"""Given the following resume, extract the experience that would be relevant to include according to the key words provided from the job posting. Think this through step by step.
    
    Key words from job posting: {key_words}
    Resume: {resume}
                \n\n"""
    output = chatbot_completition(SYS, USER)
    
    
    return output
def write_coverletter(job_desc,key_words, experience, personal_dets, instructions):
    """
    Extracts key words from a text
    """
    SYS = f"""You are an expert hiring manager and cover letter writer, responsible for tailoring a cover letter perfectly for a job posting.{instructions}  Respond using markdown"""
    
    USER = f""" Given the below job description, relevant key words and resume experiences from the candidate, complete the template. 
            Job description:{job_desc}
            Key words from job posting: {key_words}
            Relevant experiences: {experience}
            Personal details: {personal_dets}

            Template:
            Dear [Hiring Manager's Name],

            I am writing to express my interest in the [Job Title] position at [Company Name] as advertised. I am a [Your Profession] with a background in [Relevant Fields or Industries] and believe that my unique skills and experiences make me a strong candidate for this position.
            In my current/previous role as [Your Current/Previous Job Title] at [Your Current/Previous Company], I [describe a key accomplishment or responsibility that showcases relevant skills]. For example, I [specific example of a time you used a particular skill or demonstrated a particular achievement]. I believe these experiences align with the requirements you are seeking for the [Job Title] role.
            What interests me about [Company Name] is [specific details about the company that have drawn your interest]. Your commitment to [specific aspect of the company's work or culture] particularly resonates with my own professional values.
            I am particularly adept at [another skill or two that is mentioned in the job posting], as demonstrated by [another specific example]. My [mention a personal quality or qualities] will allow me to [explain how you would use your personal qualities to succeed in the job].
            Thank you for considering my application. I am excited about the possibility of contributing to [Company Name] and look forward to the opportunity to further discuss how my background, skills, and experience can benefit your team.
            Please feel free to contact me at your earliest convenience to arrange a time to discuss how I can make a meaningful contribution to your organization.

            Sincerely,

            [Your Full Name]
            """


    output = chatbot_completition(SYS, USER)
    
    
    return output
    
    
    

# Generator

def generate_cover_letter(personal_details, extra_instructions, job_descr, resume):
    try:
        key_words = extract_key_words_and_resp(job_descr)
        logger.info(f"Key words: {key_words}")
        filtered_key_words = filter_key_words(key_words,resume)
        logger.info(f"Filtered key words: {filtered_key_words}")
        relevant_exp = extract_resume_exp(filtered_key_words,resume)
        logger.info(f"Relevant experience: {relevant_exp}")
        cover_letter = write_coverletter(job_descr, filtered_key_words,relevant_exp, personal_details, extra_instructions)
        logger.info(f"Cover letter: {cover_letter}")
        file_path = create_docx(cover_letter)
        logger.info(f"File path: {file_path}")
        return file_path
    except Exception as e:
        logger.error(f"Failed to generate cover letter: {e}")
        raise

# Main app

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class GenerateDocxRequest(BaseModel):
    personal_details: str
    extra_instructions: str
    
    
# Variables to hold the values for resume and job_url
resume = None
job_descr = None

@app.post("/upload_resume_and_job_url/")
async def upload_resume_and_job_url(resume_file: UploadFile = File(...), job_link: str = Form("")):
    global resume
    global job_descr
    try:
        # Save resume file temporarily
        with open("temp_resume.docx", "wb") as buffer:
            shutil.copyfileobj(resume_file.file, buffer)
        resume_file = "temp_resume.docx"
        logging.info(f"Resume file: {resume_file}")
        resume = process_resume_file(resume_file)        
        logging.info(f"Resume: {resume}")
        job_descr = process_job_link(job_link)
        logging.info(f"Job description: {job_descr}")
        return {"status": "success"}
    except Exception as e:
        logger.error(f"Failed to upload resume and job URL: {e}")
        return {"status": "failed"}

@app.post("/generate_docx/")
async def generate_docx(request: GenerateDocxRequest):
    if resume is None or job_descr is None:
        return {"status": "failed", "message": "Resume and job URL not uploaded"}
    if not request.personal_details:
        return {"status": "failed", "message": "Personal details not provided"}
    if not request.extra_instructions:
        request.extra_instructions = ""
    file_path = generate_cover_letter(request.personal_details, request.extra_instructions, job_descr, resume)
    return FileResponse(file_path)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
    


# resume_file = None
# job_descr = None
# def process_resume_file(file_path: str):
#     # Dummy implementation that simply returns the file_path
#     return file_path

# def process_job_link(job_link: str):
#     # Dummy implementation that simply returns the job_link
#     return job_link

# def generate_cover_letter(personal_details: str, extra_instructions: str):
#     try:
#         doc = Document()
#         doc.add_heading('Personal Details', level=1)
#         doc.add_paragraph(personal_details)
#         doc.add_heading('Extra Instructions', level=1)
#         doc.add_paragraph(extra_instructions)
#         file_path = "output.docx"
#         doc.save(file_path)
#         return file_path
#     except Exception as e:
#         logger.error(f"Failed to generate cover letter: {e}")
#         raise

# @app.post("/upload_resume_and_job_url/")
# async def upload_resume_and_job_url(resume: UploadFile = File(...), job_link: str = Form("")):
#     global resume_file
#     global job_descr
#     try:
#         # Save resume file temporarily
#         with open("temp_resume.docx", "wb") as buffer:
#             shutil.copyfileobj(resume.file, buffer)
#         temp_resume_file = "temp_resume.docx"
#         resume_file = process_resume_file(temp_resume_file)
#         job_descr = process_job_link(job_link)
#         return {"status": "success"}
#     except Exception as e:
#         logger.error(f"Failed to upload resume and job URL: {e}")
#         return {"status": "failed"}

# @app.post("/generate_docx/")
# async def generate_docx(request: GenerateDocxRequest):
#     file_path = generate_cover_letter(request.personal_details, request.extra_instructions)
#     return FileResponse(file_path)




